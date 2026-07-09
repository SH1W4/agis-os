import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
from io import BytesIO
import base64
from docx import Document
import pdfkit
import tempfile

st.set_page_config(
    page_title="LogisticSmart - Relatório de Entregas",
    page_icon="📦",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.title("📦 LogisticSmart - Relatório Inteligente de Entregas")

st.markdown("### 📂 Suba seu arquivo `.xlsx` com os dados de entregas:")

uploaded_file = st.file_uploader("Escolher arquivo Excel", type=["xlsx"], help="O arquivo deve estar no formato Excel (.xlsx)")

def gerar_docx(resultado):
    doc = Document()
    doc.add_heading("Relatório de Entregas Agrupadas", level=1)
    table = doc.add_table(rows=1, cols=len(resultado.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(resultado.columns):
        hdr_cells[i].text = col
    for _, row in resultado.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    temp = BytesIO()
    doc.save(temp)
    temp.seek(0)
    return temp

def gerar_pdf(resultado):
    html = resultado.to_html(index=False)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".html") as f:
        f.write(html.encode("utf-8"))
        f.flush()
        pdf_data = pdfkit.from_file(f.name, False)
        return BytesIO(pdf_data)

if uploaded_file:
    try:
        df = pd.read_excel(uploaded_file, engine='openpyxl')

        if 'Data prevista de entrega' not in df.columns:
            st.error("❌ Coluna obrigatória 'Data prevista de entrega' não encontrada.")
            st.stop()

        df['Data prevista de entrega'] = pd.to_datetime(df['Data prevista de entrega'], errors='coerce')
        df = df.dropna(subset=['Data prevista de entrega'])

        aba = st.radio("📊 Selecione a visualização:", ["🔍 Análise de Pendentes", "✅ Entregas Realizadas"])

        hoje = datetime.now().date()
        opcoes_data = {"Hoje": hoje, "Ontem": hoje - timedelta(days=1), "Amanhã": hoje + timedelta(days=1)}
        opcao = st.selectbox("📅 Filtrar por data:", list(opcoes_data.keys()) + ["Selecionar manualmente"])

        if opcao == "Selecionar manualmente":
            data_alvo = st.date_input("📆 Escolha a data:", hoje)
        else:
            data_alvo = opcoes_data[opcao]

        df = df[df['Data prevista de entrega'].dt.date == data_alvo]

        if aba == "✅ Entregas Realizadas":
            if 'Status' in df.columns:
                df = df[df['Status'].astype(str).str.lower().str.contains("entregue")]

        colunas = df.columns

        def filtrar_opcoes(coluna):
            if coluna in colunas:
                return sorted(df[coluna].astype(str).dropna().unique().tolist())
            return []

        st.markdown("---")
        st.markdown("### 🎛️ Filtros disponíveis:")

        filtros_selecionados = {}
        for coluna in colunas:
            if df[coluna].nunique() > 1 and df[coluna].dtype == object:
                opcoes = filtrar_opcoes(coluna)
                selecionadas = st.multiselect(f"🔎 Filtrar por: {coluna}", opcoes)
                if selecionadas:
                    filtros_selecionados[coluna] = selecionadas

        for coluna, opcoes in filtros_selecionados.items():
            df = df[df[coluna].astype(str).isin(opcoes)]

        if df.empty:
            st.warning("⚠️ Nenhuma entrega encontrada para os filtros selecionados.")
        else:
            resultado = (
                df['Entregador']
                .value_counts()
                .reset_index()
                .rename(columns={'index': 'Entregador', 'Entregador': 'Quantidade'})
            )

            st.subheader("📊 Resumo por entregador")
            st.dataframe(resultado)

            buffer = BytesIO()
            resultado.to_excel(buffer, index=False, engine='openpyxl')
            buffer.seek(0)
            b64 = base64.b64encode(buffer.read()).decode()
            href = f'<a href="data:application/octet-stream;base64,{b64}" download="entregas_agrupadas_{data_alvo.strftime('%d-%m')}.xlsx">🔗 Baixar Excel</a>'
            st.markdown(href, unsafe_allow_html=True)

            st.download_button("📄 Baixar como DOCX", gerar_docx(resultado), file_name=f"relatorio_{data_alvo.strftime('%d-%m')}.docx")

            if st.button("📅 Baixar como PDF"):
                st.download_button("📅 Clique para baixar PDF", gerar_pdf(resultado), file_name=f"relatorio_{data_alvo.strftime('%d-%m')}.pdf")

    except Exception as e:
        st.error(f"❌ Erro ao processar o arquivo: {e}")
