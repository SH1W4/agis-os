"""
Utilitários para exportação de dados em diferentes formatos.
"""
import pandas as pd
from io import BytesIO
import tempfile
from datetime import datetime
from typing import Optional
import logging

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfkit
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

logger = logging.getLogger(__name__)

class ExportManager:
    """Gerenciador de exportação de dados."""
    
    def __init__(self):
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    def to_excel(self, data: pd.DataFrame, sheet_name: str = "Relatório") -> BytesIO:
        """
        Exporta dados para Excel.
        
        Args:
            data: DataFrame com os dados
            sheet_name: Nome da planilha
            
        Returns:
            Buffer com dados do Excel
        """
        buffer = BytesIO()
        
        try:
            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                # Planilha principal
                data.to_excel(writer, sheet_name=sheet_name, index=False)
                
                # Formatação básica
                workbook = writer.book
                worksheet = writer.sheets[sheet_name]
                
                # Ajustar largura das colunas
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
                
                # Adicionar metadados
                metadata_sheet = workbook.create_sheet("Informações")
                metadata_sheet['A1'] = "Relatório LogisticSmart"
                metadata_sheet['A2'] = f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
                metadata_sheet['A3'] = f"Total de registros: {len(data)}"
                
            buffer.seek(0)
            logger.info(f"Excel exportado com sucesso: {len(data)} registros")
            return buffer
            
        except Exception as e:
            logger.error(f"Erro ao exportar Excel: {e}")
            raise Exception(f"Erro na exportação Excel: {e}")
    
    def to_csv(self, data: pd.DataFrame, separator: str = ";") -> str:
        """
        Exporta dados para CSV.
        
        Args:
            data: DataFrame com os dados
            separator: Separador de campos
            
        Returns:
            String com dados CSV
        """
        try:
            csv_data = data.to_csv(
                index=False, 
                sep=separator, 
                encoding='utf-8-sig',
                date_format='%d/%m/%Y'
            )
            
            logger.info(f"CSV exportado com sucesso: {len(data)} registros")
            return csv_data
            
        except Exception as e:
            logger.error(f"Erro ao exportar CSV: {e}")
            raise Exception(f"Erro na exportação CSV: {e}")
    
    def to_docx(self, data: pd.DataFrame, title: str = "Relatório de Entregas") -> Optional[BytesIO]:
        """
        Exporta dados para Word (DOCX).
        
        Args:
            data: DataFrame com os dados
            title: Título do documento
            
        Returns:
            Buffer com dados do Word ou None se não disponível
        """
        if not DOCX_AVAILABLE:
            logger.warning("python-docx não está disponível")
            return None
        
        try:
            doc = Document()
            
            # Cabeçalho
            doc.add_heading(title, level=1)
            doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
            doc.add_paragraph(f"Total de registros: {len(data)}")
            doc.add_paragraph("")
            
            # Tabela com dados
            if not data.empty:
                table = doc.add_table(rows=1, cols=len(data.columns))
                table.style = 'Light Grid Accent 1'
                
                # Cabeçalhos
                header_cells = table.rows[0].cells
                for i, column in enumerate(data.columns):
                    header_cells[i].text = str(column)
                
                # Dados
                for _, row in data.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)
            
            # Salvar em buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            logger.info(f"DOCX exportado com sucesso: {len(data)} registros")
            return buffer
            
        except Exception as e:
            logger.error(f"Erro ao exportar DOCX: {e}")
            raise Exception(f"Erro na exportação DOCX: {e}")
    
    def to_pdf(self, data: pd.DataFrame, title: str = "Relatório de Entregas") -> Optional[BytesIO]:
        """
        Exporta dados para PDF.
        
        Args:
            data: DataFrame com os dados
            title: Título do documento
            
        Returns:
            Buffer com dados do PDF ou None se não disponível
        """
        if not PDF_AVAILABLE:
            logger.warning("pdfkit não está disponível")
            return None
        
        try:
            # Gerar HTML
            html_content = self._generate_html_report(data, title)
            
            # Configurações do PDF
            options = {
                'page-size': 'A4',
                'margin-top': '0.75in',
                'margin-right': '0.75in',
                'margin-bottom': '0.75in',
                'margin-left': '0.75in',
                'encoding': "UTF-8",
                'no-outline': None,
                'enable-local-file-access': None
            }
            
            # Gerar PDF
            pdf_data = pdfkit.from_string(html_content, False, options=options)
            
            buffer = BytesIO(pdf_data)
            buffer.seek(0)
            
            logger.info(f"PDF exportado com sucesso: {len(data)} registros")
            return buffer
            
        except Exception as e:
            logger.error(f"Erro ao exportar PDF: {e}")
            raise Exception(f"Erro na exportação PDF: {e}")
    
    def _generate_html_report(self, data: pd.DataFrame, title: str) -> str:
        """Gera HTML para o relatório PDF."""
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{title}</title>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    margin: 20px;
                    color: #333;
                }}
                
                .header {{
                    text-align: center;
                    margin-bottom: 30px;
                    border-bottom: 2px solid #08c6ff;
                    padding-bottom: 20px;
                }}
                
                .header h1 {{
                    color: #08c6ff;
                    margin: 0;
                }}
                
                .info {{
                    margin-bottom: 20px;
                    background-color: #f8f9fa;
                    padding: 15px;
                    border-radius: 5px;
                }}
                
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin-top: 20px;
                }}
                
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                }}
                
                th {{
                    background-color: #08c6ff;
                    color: white;
                    font-weight: bold;
                }}
                
                tr:nth-child(even) {{
                    background-color: #f2f2f2;
                }}
                
                .footer {{
                    margin-top: 30px;
                    text-align: center;
                    font-size: 12px;
                    color: #666;
                    border-top: 1px solid #ddd;
                    padding-top: 15px;
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>📦 {title}</h1>
                <p>LogisticSmart v2.0 - Sistema Inteligente de Análise de Entregas</p>
            </div>
            
            <div class="info">
                <p><strong>Data de Geração:</strong> {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}</p>
                <p><strong>Total de Registros:</strong> {len(data)}</p>
            </div>
            
            {data.to_html(table_id='data_table', classes='data-table', escape=False, index=False)}
            
            <div class="footer">
                <p>Relatório gerado automaticamente pelo LogisticSmart v2.0</p>
                <p>Desenvolvido por NEO-SH1W4</p>
            </div>
        </body>
        </html>
        """
        
        return html
    
    def get_available_formats(self) -> list:
        """Retorna lista de formatos disponíveis para exportação."""
        formats = ['Excel', 'CSV']
        
        if DOCX_AVAILABLE:
            formats.append('Word')
        
        if PDF_AVAILABLE:
            formats.append('PDF')
        
        return formats
    
    def export_multiple_formats(self, data: pd.DataFrame, formats: list, base_filename: str = None) -> dict:
        """
        Exporta dados em múltiplos formatos.
        
        Args:
            data: DataFrame com os dados
            formats: Lista de formatos desejados
            base_filename: Nome base do arquivo
            
        Returns:
            Dicionário com dados exportados por formato
        """
        if base_filename is None:
            base_filename = f"relatorio_logistic_{self.timestamp}"
        
        results = {}
        
        for format_name in formats:
            try:
                if format_name.lower() == 'excel':
                    results['excel'] = self.to_excel(data)
                elif format_name.lower() == 'csv':
                    results['csv'] = self.to_csv(data)
                elif format_name.lower() == 'word' and DOCX_AVAILABLE:
                    results['word'] = self.to_docx(data)
                elif format_name.lower() == 'pdf' and PDF_AVAILABLE:
                    results['pdf'] = self.to_pdf(data)
                else:
                    logger.warning(f"Formato não suportado ou não disponível: {format_name}")
                    
            except Exception as e:
                logger.error(f"Erro ao exportar formato {format_name}: {e}")
                results[format_name.lower()] = None
        
        return results

